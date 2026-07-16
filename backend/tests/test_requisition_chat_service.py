"""Conversational requisition intake — gap engine, completeness, opening
message, capture/apply, multimodal assembly, and a monkeypatched chat turn."""
import io
from datetime import datetime, timedelta, timezone

from app.llm.structured import StructuredResult
from app.models import Organization, Role
from app.platform.config import settings
from app.services import requisition_chat_service as chat
from app.services.requisition_chat_service import (
    BriefFieldChange,
    ChatAttachment,
    ChatCapture,
    apply_capture,
    build_chat_system_prompt,
    build_persisted_user_message,
    build_user_turn_content,
    compute_completeness,
    compute_gaps,
    opening_message,
    prepare_user_turn_content,
    recent_role_titles,
    run_chat_turn,
    seed_opening_message,
    warm_start_fields,
    warm_start_from_roles,
)
from app.services.requisition_intake_agent import (
    WeightedPriority,
)
from app.services.requisition_template_service import (
    DEFAULT_REQUISITION_TEMPLATE,
    resolve_template,
)
from app.services.role_brief_service import create_brief, update_brief_fields


def test_record_answer_extracts_number_from_natural_language_chip(db):
    """A quick-reply chip like "2 openings" tapped against the numeric `openings`
    field must record 2 — not 422 the whole answer. Regression for the
    "Could not record that answer" intake error."""
    org = Organization(name="Acme", slug="acme")
    db.add(org)
    db.flush()
    brief = create_brief(db, organization_id=org.id)
    template = resolve_template(org)

    chat.record_answer(db, brief, template, "openings", "2 openings")
    assert brief.openings == 2

    # A genuinely non-numeric reply still fails cleanly (no crash, no value).
    import pytest
    from fastapi import HTTPException
    with pytest.raises(HTTPException):
        chat.record_answer(db, brief, template, "openings", "lots of people")


def _brief(db, **org_kw):
    org = Organization(name="Acme", slug="acme", **org_kw)
    db.add(org)
    db.flush()
    return create_brief(db, organization_id=org.id), org


# --------------------------------------------------------------------------- #
# Gap engine + completeness
# --------------------------------------------------------------------------- #
def test_compute_gaps_lists_required_empty_in_template_order(db):
    b, _ = _brief(db)
    template = DEFAULT_REQUISITION_TEMPLATE
    gaps = compute_gaps(b, template)
    keys = [g["key"] for g in gaps]
    # All required fields, in template (section→field) order. The intake gathers
    # role substance (domain, seniority, summary, success profile, key
    # responsibilities) — not just logistics + a must-have list. Compensation is
    # HR/People's call: the agent never asks it, so it's NOT required.
    assert keys == [
        "title",
        "domain",
        "seniority",
        "summary",
        "workplace_type",
        "employment_type",
        "openings",
        "urgency",
        "must_haves",
        "success_profile",
        "responsibilities",
    ]
    # Each gap carries its section.
    title_gap = gaps[0]
    assert title_gap["section"] == "role_basics" and title_gap["label"] == "Title"


def test_gaps_shrink_as_required_fields_fill(db):
    b, _ = _brief(db)
    template = DEFAULT_REQUISITION_TEMPLATE
    update_brief_fields(db, b, title="Backend Engineer", openings=2)
    keys = [g["key"] for g in compute_gaps(b, template)]
    assert "title" not in keys and "openings" not in keys
    assert "workplace_type" in keys  # still empty


def test_empty_list_and_dict_count_as_unfilled(db):
    b, _ = _brief(db)
    update_brief_fields(db, b, must_haves=[], sourcing_signals={})
    keys = [g["key"] for g in compute_gaps(b, DEFAULT_REQUISITION_TEMPLATE)]
    assert "must_haves" in keys  # [] is empty → still a gap


def test_suggested_replies_prefers_model_supplied(db):
    b, _ = _brief(db)
    cap = ChatCapture(assistant_reply="?", suggested_replies=["A", "B", "", "C"])
    replies = chat._resolve_suggested_replies(cap, b, DEFAULT_REQUISITION_TEMPLATE)
    assert replies == ["A", "B", "C"]  # blanks stripped, model wins


def test_suggested_replies_fallback_to_select_options_of_next_gap(db):
    b, _ = _brief(db)
    template = DEFAULT_REQUISITION_TEMPLATE
    # With the role-basics required fields filled, the next required gap is
    # workplace_type (a select), so the fallback offers its options.
    update_brief_fields(
        db, b, title="Eng", seniority="Mid", summary="Builds APIs",
        custom_fields={"domain": "Banking"},
    )
    cap = ChatCapture(assistant_reply="Onsite, hybrid or remote?")  # no replies
    replies = chat._resolve_suggested_replies(cap, b, template)
    assert replies == ["Onsite", "Hybrid", "Remote"]


def test_opening_message_has_no_options_for_text_first_field(db):
    b, _ = _brief(db)
    seed_opening_message(b, DEFAULT_REQUISITION_TEMPLATE)
    # First required field (title) is text → opening offers no tappable chips.
    assert b.messages[0]["role"] == "assistant"
    assert b.messages[0]["suggested_replies"] == []


def test_completeness_math(db):
    b, _ = _brief(db)
    template = DEFAULT_REQUISITION_TEMPLATE
    assert compute_completeness(b, template) == 0
    # 11 required fields total (salary is NOT required — never asked). Fill 2 →
    # round(100*2/11) = 18.
    update_brief_fields(db, b, title="Eng", openings=1)
    assert compute_completeness(b, template) == 18
    # Fill the remaining 9 (domain / urgency / responsibilities are custom keys
    # → custom_fields; the rest are brief columns). Salary stays empty — optional.
    update_brief_fields(
        db, b,
        seniority="Mid", summary="Builds APIs",
        workplace_type="Remote", employment_type="Full-time",
        must_haves=["Python"], success_profile="Ships reliably",
        custom_fields={"urgency": "High", "domain": "Banking", "responsibilities": ["Build"]},
    )
    assert compute_completeness(b, template) == 100


def test_completeness_100_when_no_required_fields(db):
    b, _ = _brief(db)
    template = {
        "version": 1,
        "sections": [
            {"key": "s", "label": "S", "fields": [{"key": "title", "label": "T", "type": "text", "required": False}]}
        ],
    }
    assert compute_completeness(b, template) == 100


# --------------------------------------------------------------------------- #
# Opening message
# --------------------------------------------------------------------------- #
def test_opening_message_is_free_text_brief_with_checklist():
    msg = opening_message(DEFAULT_REQUISITION_TEMPLATE)
    assert msg.startswith("Hi —")
    # Free-text-first: asks for the role in their own words + names the domain.
    assert "in your own words" in msg
    assert "domain" in msg.lower()


def test_seed_opening_message_sets_single_assistant_turn(db):
    b, _ = _brief(db)
    seed_opening_message(b, DEFAULT_REQUISITION_TEMPLATE)
    assert len(b.messages) == 1
    assert b.messages[0]["role"] == "assistant"
    assert "in your own words" in b.messages[0]["content"]
    # No tappable options on the opener — free text first.
    assert b.messages[0]["suggested_replies"] == []
    assert b.messages[0]["attachments"] == []


# --------------------------------------------------------------------------- #
# Capture → apply (coercion, columns + custom_fields, no-blanking)
# --------------------------------------------------------------------------- #
def test_apply_capture_routes_columns_lists_structs_and_custom(db):
    custom_template = {
        "version": 1,
        "sections": [
            {
                "key": "role_basics",
                "label": "Role basics",
                "fields": [
                    {"key": "title", "label": "Title", "type": "text", "required": True},
                ],
            },
            {
                "key": "requirements",
                "label": "Requirements",
                "fields": [
                    {"key": "must_haves", "label": "Must-haves", "type": "list", "required": True},
                ],
            },
            {
                "key": "context",
                "label": "Context",
                "fields": [
                    {"key": "priorities", "label": "Priorities", "type": "struct_list", "required": False},
                    # An org-added custom field with NO RoleBrief column.
                    {"key": "visa_sponsorship", "label": "Visa", "type": "text", "required": False},
                ],
            },
        ],
    }
    b, _ = _brief(db)
    capture = ChatCapture(
        assistant_reply="Got it — what's the salary range?",
        open_questions=["salary range?"],
        title="Senior Backend Engineer",
        must_haves=["Python", "Postgres"],
        priorities=[WeightedPriority(factor="domain", weight="high")],
        custom={"visa_sponsorship": "Yes, we sponsor"},
    )
    apply_capture(db, b, capture, custom_template)

    # Column (text), list, struct_list applied to the right columns.
    assert b.title == "Senior Backend Engineer"
    assert b.must_haves == ["Python", "Postgres"]
    assert b.priorities == [{"factor": "domain", "weight": "high"}]
    # Custom (no-column) key landed in custom_fields.
    assert b.custom_fields["visa_sponsorship"] == "Yes, we sponsor"
    # open_questions persisted to agent_state.
    assert b.agent_state["open_questions"] == ["salary range?"]
    # completeness: 2/2 required filled → 100.
    assert b.completeness == 100


def test_apply_capture_does_not_blank_previously_captured(db):
    b, _ = _brief(db)
    update_brief_fields(db, b, title="Original Title", must_haves=["Python"])
    # A turn that captures nothing for title/must_haves must not wipe them.
    capture = ChatCapture(assistant_reply="Thanks!", department="Platform")
    apply_capture(db, b, capture, DEFAULT_REQUISITION_TEMPLATE)
    assert b.title == "Original Title"
    assert b.must_haves == ["Python"]
    assert b.department == "Platform"


def test_apply_capture_number_coercion(db):
    b, _ = _brief(db)
    # Model may emit a string for a number-typed field; coerce it.
    capture = ChatCapture(assistant_reply="ok", openings="3", salary_min=120000)
    apply_capture(db, b, capture, DEFAULT_REQUISITION_TEMPLATE)
    assert b.openings == 3
    assert b.salary_min == 120000


def test_apply_capture_target_start_date_maps_to_target_start_column(db):
    b, _ = _brief(db)
    capture = ChatCapture(assistant_reply="ok", target_start_date="2026-09-01")
    apply_capture(db, b, capture, DEFAULT_REQUISITION_TEMPLATE)
    assert b.target_start == "2026-09-01"


def test_apply_capture_sourcing_signals_list_and_process_text(db):
    # sourcing_signals (list) and process (longtext) match the template's
    # declared types — stored as a cleaned list of strings and a plain string.
    b, _ = _brief(db)
    capture = ChatCapture(
        assistant_reply="ok",
        sourcing_signals=["Ex-Stripe", "Strong OSS presence"],
        process="3 rounds: screen, take-home, onsite. Urgent.",
    )
    apply_capture(db, b, capture, DEFAULT_REQUISITION_TEMPLATE)
    assert b.sourcing_signals == ["Ex-Stripe", "Strong OSS presence"]
    assert b.process == "3 rounds: screen, take-home, onsite. Urgent."


def test_apply_capture_routes_builtin_template_only_fields_to_custom_fields(db):
    """Core default-template fields must be first-class capture properties even
    though RoleBrief stores them in custom_fields."""
    b, _ = _brief(db)
    capture = ChatCapture(
        assistant_reply="Captured from the spec.",
        domain="Regulated banking",
        urgency="Urgent",
        benefits=["Private medical", "Learning budget"],
        responsibilities=["Build AI services", "Own production reliability"],
    )
    apply_capture(db, b, capture, DEFAULT_REQUISITION_TEMPLATE)

    assert b.custom_fields["domain"] == "Regulated banking"
    assert b.custom_fields["urgency"] == "Urgent"
    assert b.custom_fields["benefits"] == ["Private medical", "Learning budget"]
    assert b.custom_fields["responsibilities"] == [
        "Build AI services",
        "Own production reliability",
    ]
    remaining = {gap["key"] for gap in compute_gaps(b, DEFAULT_REQUISITION_TEMPLATE)}
    assert "domain" not in remaining
    assert "urgency" not in remaining
    assert "responsibilities" not in remaining


def test_apply_capture_amends_lists_item_by_item_and_revises_canonical_spec(db):
    b, _ = _brief(db)
    update_brief_fields(
        db,
        b,
        must_haves=["Python", "Java"],
        agent_state={
            "jd_override": "# Engineer\n\nMust have Python and Java.",
            "canonical_spec_mode": "verbatim",
            "job_spec_revision": 1,
        },
    )
    capture = ChatCapture(
        assistant_reply="Updated.",
        change_mode="amend",
        changes=[
            BriefFieldChange(key="must_haves", operation="remove", value=["Java"]),
            BriefFieldChange(
                key="must_haves", operation="add", value=["Azure", "azure"]
            ),
        ],
        canonical_job_spec="# Engineer\n\nMust have Python and Azure.",
    )

    apply_capture(db, b, capture, DEFAULT_REQUISITION_TEMPLATE)

    assert b.must_haves == ["Python", "Azure"]
    assert b.agent_state["jd_override"].endswith("Python and Azure.")
    assert b.agent_state["job_spec_revision"] == 2
    assert b.agent_state["job_spec_last_change_mode"] == "amend"


def test_apply_capture_replace_resets_role_content_but_preserves_related_link(db):
    b, _ = _brief(db)
    source = Role(organization_id=b.organization_id, name="Original ATS role")
    db.add(source)
    db.flush()
    b.source_role_id = source.id
    update_brief_fields(
        db,
        b,
        title="Old Engineer",
        summary="Old summary",
        must_haves=["Java"],
        preferred=["Kafka"],
        custom_fields={"domain": "Banking", "private_note": "keep"},
        agent_state={"jd_override": "OLD JD", "job_spec_revision": 1},
    )
    capture = ChatCapture(
        assistant_reply="Replaced.",
        change_mode="replace",
        title="AI Engineer",
        must_haves=["Python"],
        domain="Healthcare",
        canonical_job_spec="# AI Engineer\n\nMust have Python.",
    )

    apply_capture(db, b, capture, DEFAULT_REQUISITION_TEMPLATE)

    assert b.source_role_id == source.id
    assert b.title == "AI Engineer"
    assert b.summary is None
    assert b.must_haves == ["Python"]
    assert b.preferred == []
    assert b.custom_fields == {"domain": "Healthcare", "private_note": "keep"}
    assert b.agent_state["jd_override"].startswith("# AI Engineer")
    assert b.agent_state["job_spec_last_change_mode"] == "replace"


def test_apply_capture_clarification_is_non_mutating_and_keeps_pending_spec(db):
    b, _ = _brief(db)
    update_brief_fields(
        db,
        b,
        title="Current role",
        must_haves=["Python"],
        agent_state={"jd_override": "CURRENT JD", "job_spec_revision": 1},
    )
    capture = ChatCapture(
        assistant_reply="Replace or amend?",
        change_mode="clarify",
        title="Should not apply",
        changes=[BriefFieldChange(key="must_haves", operation="clear")],
        pending_job_spec="PROPOSED JD",
    )

    apply_capture(db, b, capture, DEFAULT_REQUISITION_TEMPLATE)

    assert b.title == "Current role"
    assert b.must_haves == ["Python"]
    assert b.agent_state["jd_override"] == "CURRENT JD"
    assert b.agent_state["pending_job_spec_source"] == "PROPOSED JD"


def test_apply_capture_never_leaves_a_stale_verbatim_spec_after_field_edit(db):
    b, _ = _brief(db)
    update_brief_fields(
        db,
        b,
        title="Old title",
        agent_state={"jd_override": "OLD JD", "job_spec_revision": 1},
    )

    apply_capture(
        db,
        b,
        ChatCapture(assistant_reply="Updated.", title="New title"),
        DEFAULT_REQUISITION_TEMPLATE,
    )

    assert b.title == "New title"
    assert "jd_override" not in b.agent_state
    assert b.agent_state["canonical_spec_mode"] == "structured"
    assert b.agent_state["job_spec_revision"] == 2


# --------------------------------------------------------------------------- #
# Multimodal assembly
# --------------------------------------------------------------------------- #
def test_persisted_user_message_records_attachment_metadata():
    msg = build_persisted_user_message(
        "see attached",
        [
            ChatAttachment(name="call.vtt", content_type="text/vtt", content=b"x"),
            ChatAttachment(name="board.png", content_type="image/png", content=b"y"),
            ChatAttachment(name="jd.pdf", content_type="application/pdf", content=b"z"),
        ],
    )
    assert msg["role"] == "user" and msg["content"] == "see attached"
    kinds = {a["name"]: a["kind"] for a in msg["attachments"]}
    assert kinds == {"call.vtt": "transcript", "board.png": "image", "jd.pdf": "file"}


def test_transcript_text_reaches_user_turn_content():
    content = build_user_turn_content(
        "here are my notes",
        [ChatAttachment(name="kickoff.txt", content_type="text/plain", content=b"We need a staff PM")],
    )
    # No image → plain string content.
    assert isinstance(content, str)
    assert "here are my notes" in content
    assert "[Attached transcript: kickoff.txt]" in content
    assert "We need a staff PM" in content


def test_transcript_and_filename_control_characters_are_sanitized():
    attachment = ChatAttachment(
        name="job\x00spec.txt",
        content_type="text/plain",
        content=b"Build APIs\x00Own reliability",
    )
    persisted = build_persisted_user_message("Use\x00this", [attachment])
    content, source = prepare_user_turn_content("Use\x00this", [attachment])

    assert "\x00" not in persisted["content"]
    assert "\x00" not in persisted["attachments"][0]["name"]
    assert "\x00" not in content
    assert "\x00" not in source
    assert "Build APIsOwn reliability" in source


def test_large_text_attachment_is_bounded_before_model_and_persistence():
    large_text = ("TITLE AND SUMMARY\n" + ("x" * 120_000) + "\nREQUIREMENTS END").encode()
    content, source = prepare_user_turn_content(
        "Use this",
        [
            ChatAttachment(
                name="large.txt",
                content_type="text/plain",
                content=large_text,
            )
        ],
    )

    assert len(source) < 31_000
    assert len(content) < 31_100
    assert "TITLE AND SUMMARY" in source
    assert "REQUIREMENTS END" in source
    assert "content truncated for safe processing" in source


def test_image_attachment_produces_base64_image_block():
    content = build_user_turn_content(
        "what does this say",
        [ChatAttachment(name="shot.png", content_type="image/png", content=b"\x89PNG-bytes")],
    )
    # Image present → list of content blocks, with one image block (base64).
    assert isinstance(content, list)
    image_blocks = [b for b in content if b.get("type") == "image"]
    assert len(image_blocks) == 1
    src = image_blocks[0]["source"]
    assert src["type"] == "base64" and src["media_type"] == "image/png" and src["data"]
    # The text part is preserved alongside the image.
    assert any(b.get("type") == "text" and "what does this say" in b.get("text", "") for b in content)


def test_image_extension_is_recognised_with_octet_stream_mime():
    attachment = ChatAttachment(
        name="SPEC.PNG",
        content_type="application/octet-stream",
        content=b"\x89PNG-bytes",
    )
    persisted = build_persisted_user_message("", [attachment])
    assert persisted["attachments"] == [{"name": "SPEC.PNG", "kind": "image"}]

    content = build_user_turn_content("read this", [attachment])
    image_block = next(block for block in content if block.get("type") == "image")
    assert image_block["source"]["media_type"] == "image/png"


def test_docx_attachment_is_extracted_once_and_returned_as_recoverable_source(
    monkeypatch,
):
    calls = []

    def fake_extract_text(content, extension):
        calls.append((content, extension))
        return "Senior AI Engineer\nBuild production RAG systems"

    monkeypatch.setattr("app.services.document_service.extract_text", fake_extract_text)
    content, source = prepare_user_turn_content(
        "full job spec attached",
        [
            ChatAttachment(
                name="role.docx",
                content_type=(
                    "application/vnd.openxmlformats-officedocument."
                    "wordprocessingml.document"
                ),
                content=b"docx-bytes",
            )
        ],
    )

    assert calls == [(b"docx-bytes", "docx")]
    assert "[Attached document: role.docx]" in content
    assert "Build production RAG systems" in content
    assert source == content.split("\n\n", 1)[1]


def test_real_docx_table_content_reaches_recoverable_source():
    from docx import Document

    document = Document()
    document.add_heading("Senior AI Engineer", level=1)
    table = document.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Key responsibilities"
    table.cell(0, 1).text = "Build production RAG services\nOwn model reliability"
    table.cell(1, 0).text = "Domain"
    table.cell(1, 1).text = "Regulated banking"
    stream = io.BytesIO()
    document.save(stream)

    content, source = prepare_user_turn_content(
        "Use this full job spec",
        [
            ChatAttachment(
                name="table-spec.docx",
                content_type=(
                    "application/vnd.openxmlformats-officedocument."
                    "wordprocessingml.document"
                ),
                content=stream.getvalue(),
            )
        ],
    )

    assert "Senior AI Engineer" in content
    assert "Key responsibilities" in source
    assert "Build production RAG services" in source
    assert "Own model reliability" in source
    assert "Domain" in source
    assert "Regulated banking" in source


# --------------------------------------------------------------------------- #
# Orchestrated chat turn (LLM monkeypatched — no Anthropic)
# --------------------------------------------------------------------------- #
def test_run_chat_turn_applies_capture_appends_messages_shrinks_gaps(db, monkeypatch):
    b, _org = _brief(db)
    seed_opening_message(b, resolve_template(_org))
    db.flush()

    captured_calls = {}

    def fake_generate_structured(client, **kwargs):
        # Record the constructed LLM input so we can assert on it.
        captured_calls["messages"] = kwargs["messages"]
        captured_calls["system"] = kwargs["system"]
        captured_calls["feature"] = kwargs["metering"].feature
        value = ChatCapture(
            assistant_reply="Great — onsite or remote, and how many openings?",
            open_questions=["workplace_type?", "openings?"],
            title="Backend Engineer",
            must_haves=["Python", "Postgres"],
            salary_min=100000,
            salary_max=140000,
            salary_currency="USD",
        )
        return StructuredResult(value=value, ok=True)

    monkeypatch.setattr(chat, "generate_structured", fake_generate_structured)

    result = run_chat_turn(
        db, b,
        message="We need a backend engineer, Python + Postgres, $100-140k.",
        attachments=[ChatAttachment(name="notes.txt", content_type="text/plain", content=b"team is small")],
        client=object(),
        model="test-model",
    )
    assert result.ok

    # Captured values applied to the right columns.
    assert b.title == "Backend Engineer"
    assert b.must_haves == ["Python", "Postgres"]
    assert b.salary_currency == "USD"
    # source_kind defaulted to conversational.
    assert b.source_kind == "conversational"
    # open_questions persisted.
    assert b.agent_state["open_questions"] == ["workplace_type?", "openings?"]

    # Transcript: opening + user + assistant reply.
    roles = [m["role"] for m in b.messages]
    assert roles == ["assistant", "user", "assistant"]
    assert b.messages[1]["content"].startswith("We need a backend engineer")
    assert b.messages[1]["attachments"] == [{"name": "notes.txt", "kind": "transcript"}]
    # Attachment replies are grounded in POST-capture gaps, not the model's
    # stale pre-capture question. Title was captured, so domain is now first.
    assert b.messages[2]["content"].startswith("I've amended the draft.")
    assert "Updated: Must-haves" in b.messages[2]["content"]
    assert "What domain or industry" in b.messages[2]["content"]

    # The transcript file content reached the LLM input (multimodal assertion).
    last_user = captured_calls["messages"][-1]
    assert last_user["role"] == "user"
    assert "[Attached transcript: notes.txt]" in last_user["content"]
    assert "team is small" in last_user["content"]
    # Decoded source survives the turn and is also supplied as source data in
    # the system prompt, with all gaps exposed for exhaustive extraction.
    assert "[Attached transcript: notes.txt]" in b.raw_input
    assert "team is small" in b.raw_input
    assert "RECOVERABLE SOURCE MATERIAL" in captured_calls["system"]
    assert "EXTRACT EXHAUSTIVELY" in captured_calls["system"]
    # Metered under the right feature.
    assert captured_calls["feature"] == "requisition_intake_chat"

    # Gaps shrank: title/salary/must_haves now filled; workplace_type/openings remain.
    keys = [g["key"] for g in compute_gaps(b, resolve_template(_org))]
    assert "title" not in keys and "must_haves" not in keys
    assert "workplace_type" in keys and "openings" in keys
    # completeness recomputed: 2/11 required filled (title + must_haves; salary
    # is captured but NOT required, so it doesn't count toward completeness).
    assert b.completeness == round(100 * 2 / 11)


def test_attachment_source_is_available_to_later_chat_turns(db, monkeypatch):
    b, org = _brief(db)
    template = resolve_template(org)
    systems = []

    def fake_generate_structured(client, **kwargs):
        systems.append(kwargs["system"])
        if len(systems) == 1:
            return StructuredResult(
                value=ChatCapture(
                    assistant_reply="What is the title?",
                    title="AI Engineer",
                    domain="Banking",
                ),
                ok=True,
            )
        return StructuredResult(
            value=ChatCapture(assistant_reply="What seniority level?"), ok=True
        )

    monkeypatch.setattr(chat, "generate_structured", fake_generate_structured)
    run_chat_turn(
        db,
        b,
        message="Please use this spec",
        attachments=[
            ChatAttachment(
                name="job-spec.txt",
                content_type="text/plain",
                content=b"Own our production RAG platform in regulated banking.",
            )
        ],
        template=template,
        client=object(),
        model="m",
    )
    run_chat_turn(
        db,
        b,
        message="Continue",
        template=template,
        client=object(),
        model="m",
    )

    assert "Own our production RAG platform" in b.raw_input
    assert "Own our production RAG platform" in systems[1]
    assert "RECOVERABLE SOURCE MATERIAL" in systems[1]


def test_source_remains_in_exhaustive_mode_until_capture_changes_brief(
    db, monkeypatch
):
    b, org = _brief(db)
    template = resolve_template(org)
    systems = []

    def fake_generate_structured(client, **kwargs):
        systems.append(kwargs["system"])
        capture = (
            ChatCapture(assistant_reply="I couldn't find a field yet.")
            if len(systems) < 3
            else ChatCapture(assistant_reply="Captured.", title="AI Engineer")
        )
        return StructuredResult(value=capture, ok=True)

    monkeypatch.setattr(chat, "generate_structured", fake_generate_structured)
    first = run_chat_turn(
        db,
        b,
        message="Use this",
        attachments=[
            ChatAttachment(
                name="spec.txt",
                content_type="text/plain",
                content=b"AI engineering role",
            )
        ],
        template=template,
        client=object(),
        model="m",
    )
    run_chat_turn(db, b, message="Try again", template=template, client=object(), model="m")
    run_chat_turn(db, b, message="AI Engineer", template=template, client=object(), model="m")
    run_chat_turn(db, b, message="Continue", template=template, client=object(), model="m")

    assert "EXTRACT EXHAUSTIVELY" in systems[0]
    assert "EXTRACT EXHAUSTIVELY" in systems[1]
    assert "EXTRACT EXHAUSTIVELY" in systems[2]
    assert "EXTRACT EXHAUSTIVELY" not in systems[3]
    assert "didn't add any new brief fields" in first.value.assistant_reply
    assert "I've populated" not in first.value.assistant_reply


def test_legacy_related_role_override_gets_one_exhaustive_hydration_turn(
    db, monkeypatch
):
    b, org = _brief(db)
    source = Role(
        organization_id=org.id,
        name="AI Engineer",
        source="workable",
        workable_job_id="legacy-source",
    )
    db.add(source)
    db.flush()
    b.source_role_id = source.id
    b.agent_state = {
        "jd_override": "# AI Engineer\n\n## Responsibilities\n- Own production RAG reliability."
    }
    systems = []

    def fake_generate_structured(client, **kwargs):
        systems.append(kwargs["system"])
        return StructuredResult(
            value=ChatCapture(
                assistant_reply="What is the title?",
                title="AI Engineer",
                responsibilities=["Own production RAG reliability."],
            ),
            ok=True,
        )

    monkeypatch.setattr(chat, "generate_structured", fake_generate_structured)
    run_chat_turn(
        db,
        b,
        message="Continue",
        template=resolve_template(org),
        client=object(),
        model="m",
    )
    run_chat_turn(
        db,
        b,
        message="Continue again",
        template=resolve_template(org),
        client=object(),
        model="m",
    )

    assert "Own production RAG reliability" in systems[0]
    assert "EXTRACT EXHAUSTIVELY" in systems[0]
    assert "EXTRACT EXHAUSTIVELY" not in systems[1]
    assert b.raw_input is None  # chat working-copy hydration owns persistence
    assert b.custom_fields["responsibilities"] == [
        "Own production RAG reliability."
    ]


def test_current_job_spec_is_the_only_active_source_not_old_provenance(db, monkeypatch):
    b, org = _brief(db)
    b.raw_input = "Original attached source"
    b.agent_state = {"jd_override": "Current edited job specification"}
    seen = {}

    def fake_generate_structured(client, **kwargs):
        seen["system"] = kwargs["system"]
        return StructuredResult(value=ChatCapture(assistant_reply="Next?"), ok=True)

    monkeypatch.setattr(chat, "generate_structured", fake_generate_structured)
    run_chat_turn(
        db,
        b,
        message="Continue",
        template=resolve_template(org),
        client=object(),
        model="m",
    )

    assert "Original attached source" not in seen["system"]
    assert "Current edited job specification" in seen["system"]
    assert "[ACTIVE CANONICAL JOB SPEC]" in seen["system"]


def test_related_role_full_replacement_changes_only_draft_and_canonical_spec(
    db, monkeypatch
):
    b, org = _brief(db)
    source = Role(
        organization_id=org.id,
        name="Original Java Engineer",
        job_spec_text="# Original\n\nJava role",
    )
    db.add(source)
    db.flush()
    b.source_role_id = source.id
    b.title = "Original Java Engineer · Related"
    b.must_haves = ["Java"]
    b.preferred = ["Kafka"]
    b.raw_input = source.job_spec_text
    b.agent_state = {
        "jd_override": source.job_spec_text,
        "canonical_spec_mode": "verbatim",
        "job_spec_revision": 1,
    }
    new_spec = "# AI Engineer\n\nBuild RAG systems with Python."
    monkeypatch.setattr(
        chat,
        "generate_structured",
        lambda *args, **kwargs: StructuredResult(
            value=ChatCapture(
                assistant_reply="I used the new spec.",
                change_mode="replace",
                title="AI Engineer",
                must_haves=["Python"],
                domain="Banking",
                responsibilities=["Build production RAG systems"],
                canonical_job_spec=new_spec,
            ),
            ok=True,
        ),
    )

    result = run_chat_turn(
        db,
        b,
        message="Use this new job spec instead",
        attachments=[
            ChatAttachment(
                name="new-role.txt",
                content_type="text/plain",
                content=new_spec.encode(),
            )
        ],
        template=resolve_template(org),
        client=object(),
        model="m",
    )

    assert b.source_role_id == source.id
    assert source.name == "Original Java Engineer"
    assert source.job_spec_text.endswith("Java role")
    assert b.title == "AI Engineer"
    assert b.must_haves == ["Python"]
    assert b.preferred == []
    assert b.agent_state["jd_override"] == new_spec
    assert b.agent_state["job_spec_revision"] == 2
    assert "# Original" in b.raw_input
    assert "Build RAG systems with Python" in b.raw_input
    assert result.value.assistant_reply.startswith(
        "I've replaced the role content in this related-role draft."
    )
    assert "original ATS role and shared candidate pool are unchanged" in (
        result.value.assistant_reply
    )


def test_related_role_requirement_refinement_uses_semantic_operations(db, monkeypatch):
    b, org = _brief(db)
    source = Role(organization_id=org.id, name="Original role")
    db.add(source)
    db.flush()
    b.source_role_id = source.id
    b.must_haves = ["Python", "Java"]
    b.agent_state = {
        "jd_override": "Must have Python and Java.",
        "canonical_spec_mode": "verbatim",
        "job_spec_revision": 1,
    }
    monkeypatch.setattr(
        chat,
        "generate_structured",
        lambda *args, **kwargs: StructuredResult(
            value=ChatCapture(
                assistant_reply="Updated the requirements.",
                change_mode="amend",
                changes=[
                    BriefFieldChange(
                        key="must_haves", operation="remove", value=["Java"]
                    ),
                    BriefFieldChange(
                        key="must_haves", operation="add", value=["Azure"]
                    ),
                ],
                canonical_job_spec="Must have Python and Azure.",
            ),
            ok=True,
        ),
    )

    result = run_chat_turn(
        db,
        b,
        message="Keep everything else, remove Java and add Azure",
        template=resolve_template(org),
        client=object(),
        model="m",
    )

    assert b.must_haves == ["Python", "Azure"]
    assert b.agent_state["jd_override"] == "Must have Python and Azure."
    assert result.value.assistant_reply.startswith(
        "I've amended this related-role draft"
    )
    assert "Updated: Must-haves" in result.value.assistant_reply


def test_ambiguous_full_spec_waits_for_replace_or_amend_choice(db, monkeypatch):
    b, org = _brief(db)
    b.title = "Current role"
    b.must_haves = ["Python"]
    b.agent_state = {
        "jd_override": "CURRENT JD",
        "canonical_spec_mode": "verbatim",
        "job_spec_revision": 1,
    }
    monkeypatch.setattr(
        chat,
        "generate_structured",
        lambda *args, **kwargs: StructuredResult(
            value=ChatCapture(
                assistant_reply="I need to know how to apply this.",
                change_mode="clarify",
                title="Must not apply yet",
            ),
            ok=True,
        ),
    )

    result = run_chat_turn(
        db,
        b,
        message="Here is another spec",
        attachments=[
            ChatAttachment(
                name="proposal.txt",
                content_type="text/plain",
                content=b"# Different role\n\nMust have Go.",
            )
        ],
        template=resolve_template(org),
        client=object(),
        model="m",
    )

    assert b.title == "Current role"
    assert b.must_haves == ["Python"]
    assert b.agent_state["jd_override"] == "CURRENT JD"
    assert "Must have Go" in b.agent_state["pending_job_spec_source"]
    assert result.value.suggested_replies == [
        "Replace current draft",
        "Apply differences only",
    ]
    assert "Should I replace the current draft" in result.value.assistant_reply
    assert b.messages[-1]["change_mode"] == "clarify"


def test_new_role_full_spec_becomes_canonical_then_refinements_amend_it(
    db, monkeypatch
):
    b, org = _brief(db)
    captures = iter(
        [
            ChatCapture(
                assistant_reply="Captured the spec.",
                change_mode="replace",
                title="Backend Engineer",
                must_haves=["Python"],
                canonical_job_spec="# Backend Engineer\n\nMust have Python.",
            ),
            ChatCapture(
                assistant_reply="Added Postgres.",
                change_mode="amend",
                changes=[
                    BriefFieldChange(
                        key="must_haves", operation="add", value=["Postgres"]
                    )
                ],
                canonical_job_spec=(
                    "# Backend Engineer\n\nMust have Python and Postgres."
                ),
            ),
        ]
    )
    monkeypatch.setattr(
        chat,
        "generate_structured",
        lambda *args, **kwargs: StructuredResult(value=next(captures), ok=True),
    )

    run_chat_turn(
        db,
        b,
        message="Use this full job spec",
        attachments=[
            ChatAttachment(
                name="job.txt",
                content_type="text/plain",
                content=b"# Backend Engineer\n\nMust have Python.",
            )
        ],
        template=resolve_template(org),
        client=object(),
        model="m",
    )
    refined = run_chat_turn(
        db,
        b,
        message="Keep everything else and add Postgres as a must-have",
        template=resolve_template(org),
        client=object(),
        model="m",
    )

    assert b.must_haves == ["Python", "Postgres"]
    assert b.agent_state["jd_override"].endswith("Python and Postgres.")
    assert b.agent_state["job_spec_revision"] == 2
    assert refined.value.assistant_reply.startswith("I've amended the draft.")
    assert "Updated: Must-haves" in refined.value.assistant_reply


def test_existing_and_related_draft_changes_use_reasoning_model(db, monkeypatch):
    first, org = _brief(db)
    related = create_brief(db, organization_id=org.id)
    source = Role(organization_id=org.id, name="Source role")
    db.add(source)
    db.flush()
    related.source_role_id = source.id
    related.agent_state = {"jd_override": "CURRENT JD"}
    seen_models = []
    monkeypatch.setattr(settings, "CLAUDE_CHAT_MODEL", "haiku-test")
    monkeypatch.setattr(settings, "CLAUDE_MODEL", "sonnet-test")

    def fake_generate_structured(client, **kwargs):
        seen_models.append(kwargs["model"])
        return StructuredResult(value=ChatCapture(assistant_reply="Next?"), ok=True)

    monkeypatch.setattr(chat, "generate_structured", fake_generate_structured)
    run_chat_turn(
        db,
        first,
        message="We need an engineer",
        template=resolve_template(org),
        client=object(),
    )
    run_chat_turn(
        db,
        related,
        message="Change the requirements",
        template=resolve_template(org),
        client=object(),
    )

    assert seen_models == ["haiku-test", "sonnet-test"]


def test_corrupt_document_does_not_claim_the_brief_was_populated(db, monkeypatch):
    b, org = _brief(db)
    monkeypatch.setattr(
        "app.services.document_service.extract_text",
        lambda *_args, **_kwargs: (_ for _ in ()).throw(ValueError("corrupt")),
    )
    monkeypatch.setattr(
        chat,
        "generate_structured",
        lambda *args, **kwargs: StructuredResult(
            value=ChatCapture(
                assistant_reply="I couldn't extract text from that document."
            ),
            ok=True,
        ),
    )

    result = run_chat_turn(
        db,
        b,
        message="Use this",
        attachments=[
            ChatAttachment(
                name="broken.pdf",
                content_type="application/pdf",
                content=b"not-a-pdf",
            )
        ],
        template=resolve_template(org),
        client=object(),
        model="m",
    )

    assert "couldn't extract usable content" in result.value.assistant_reply
    assert "I've populated the brief" not in result.value.assistant_reply


def test_partial_attachment_failure_is_disclosed_after_other_source_is_captured(
    db, monkeypatch
):
    b, org = _brief(db)
    monkeypatch.setattr(
        "app.services.document_service.extract_text",
        lambda *_args, **_kwargs: (_ for _ in ()).throw(ValueError("corrupt")),
    )
    monkeypatch.setattr(
        chat,
        "generate_structured",
        lambda *args, **kwargs: StructuredResult(
            value=ChatCapture(assistant_reply="Captured.", title="AI Engineer"),
            ok=True,
        ),
    )

    result = run_chat_turn(
        db,
        b,
        message="Use both files",
        attachments=[
            ChatAttachment(
                name="notes.txt",
                content_type="text/plain",
                content=b"AI Engineer",
            ),
            ChatAttachment(
                name="broken.pdf",
                content_type="application/pdf",
                content=b"not-a-pdf",
            ),
        ],
        template=resolve_template(org),
        client=object(),
        model="m",
    )

    assert result.value.assistant_reply.startswith("I've amended the draft")
    assert "couldn't read every attachment" in result.value.assistant_reply


def test_chat_replaces_false_publish_claim_with_capability_grounding(
    db, monkeypatch
):
    b, org = _brief(db)

    monkeypatch.setattr(
        chat,
        "generate_structured",
        lambda *args, **kwargs: StructuredResult(
            value=ChatCapture(
                assistant_reply=(
                    "Perfect — the job is published and live for sourcing now."
                )
            ),
            ok=True,
        ),
    )
    result = run_chat_turn(
        db,
        b,
        message="Yes, publish now",
        template=resolve_template(org),
        client=object(),
        model="m",
    )

    reply = result.value.assistant_reply
    assert "published and live" not in reply
    assert "haven't performed the requested action" in reply
    assert "chat only saves brief fields" in reply
    assert b.status == "draft"
    assert b.role_id is None


def test_action_receipts_are_grounded_without_domain_language_false_positives(
    db, monkeypatch
):
    b, org = _brief(db)
    replies = iter(
        [
            "Done — your job is live.",
            "Got it — Publishing Manager is the title.",
            "Got it — publishing weekly reports is a responsibility.",
            "Done — I have posted the job.",
            "Done — the opening is active.",
            "Candidates should have published research.",
        ]
    )
    monkeypatch.setattr(
        chat,
        "generate_structured",
        lambda *args, **kwargs: StructuredResult(
            value=ChatCapture(assistant_reply=next(replies)), ok=True
        ),
    )

    launched = run_chat_turn(
        db,
        b,
        message="Please launch it",
        template=resolve_template(org),
        client=object(),
        model="m",
    )
    titled = run_chat_turn(
        db,
        b,
        message="We need a publishing manager",
        template=resolve_template(org),
        client=object(),
        model="m",
    )
    responsibility = run_chat_turn(
        db,
        b,
        message="They need to publish weekly reports",
        template=resolve_template(org),
        client=object(),
        model="m",
    )
    posted = run_chat_turn(
        db,
        b,
        message="Please post it",
        template=resolve_template(org),
        client=object(),
        model="m",
    )
    activated = run_chat_turn(
        db,
        b,
        message="Activate the opening",
        template=resolve_template(org),
        client=object(),
        model="m",
    )
    research = run_chat_turn(
        db,
        b,
        message="The role will publish research",
        template=resolve_template(org),
        client=object(),
        model="m",
    )

    assert "job is live" not in launched.value.assistant_reply
    assert "haven't performed the requested action" in launched.value.assistant_reply
    assert titled.value.assistant_reply == "Got it — Publishing Manager is the title."
    assert responsibility.value.assistant_reply == (
        "Got it — publishing weekly reports is a responsibility."
    )
    for result in (posted, activated):
        assert "haven't performed the requested action" in result.value.assistant_reply
    assert research.value.assistant_reply == "Candidates should have published research."


def test_document_turn_asks_only_first_post_capture_gap(db, monkeypatch):
    b, org = _brief(db)
    template = resolve_template(org)
    monkeypatch.setattr(
        chat,
        "generate_structured",
        lambda *args, **kwargs: StructuredResult(
            value=ChatCapture(
                assistant_reply="What role are you hiring for?",
                title="AI Engineer",
                domain="Banking",
                seniority="Senior",
                summary="Build production AI applications",
            ),
            ok=True,
        ),
    )

    result = run_chat_turn(
        db,
        b,
        message="",
        attachments=[
            ChatAttachment(
                name="role.txt",
                content_type="text/plain",
                content=b"Senior AI Engineer in banking",
            )
        ],
        template=template,
        client=object(),
        model="m",
    )

    assert "What role are you hiring for?" not in result.value.assistant_reply
    assert "Is this onsite, hybrid, or remote?" in result.value.assistant_reply


def test_run_chat_turn_image_block_reaches_llm(db, monkeypatch):
    b, _org = _brief(db)

    seen = {}

    def fake_generate_structured(client, **kwargs):
        seen["messages"] = kwargs["messages"]
        return StructuredResult(
            value=ChatCapture(assistant_reply="I can see the whiteboard."), ok=True
        )

    monkeypatch.setattr(chat, "generate_structured", fake_generate_structured)
    run_chat_turn(
        db, b,
        message="here's the whiteboard from our kickoff",
        attachments=[ChatAttachment(name="wb.png", content_type="image/png", content=b"\x89PNGdata")],
        client=object(),
        model="test-model",
    )
    last_user = seen["messages"][-1]
    assert isinstance(last_user["content"], list)
    assert any(blk.get("type") == "image" for blk in last_user["content"])


def test_run_chat_turn_failure_rolls_back_capture_but_keeps_user_message(db, monkeypatch):
    b, _org = _brief(db)
    monkeypatch.setattr(
        chat, "generate_structured",
        lambda *a, **k: StructuredResult(value=None, ok=False, error_reason="boom"),
    )
    result = run_chat_turn(db, b, message="a backend engineer", client=object(), model="m")
    assert not result.ok
    # Nothing captured.
    assert b.title is None
    # The user message was still appended (the turn happened).
    assert b.messages[-1]["role"] == "user"


# --------------------------------------------------------------------------- #
# Warm-start: recency-biased prefill + recent-roles context
# --------------------------------------------------------------------------- #
def _org(db, **org_kw):
    org = Organization(name="Acme", slug="acme", **org_kw)
    db.add(org)
    db.flush()
    return org


def test_warm_start_fields_takes_most_recent_non_empty_per_field(db):
    org = _org(db)
    # Oldest brief: has a city + workplace_type.
    older = create_brief(db, organization_id=org.id)
    update_brief_fields(
        db, older,
        location_city="Abu Dhabi", workplace_type="Onsite", department="Eng",
    )
    # Newer brief: a different city, no workplace_type, no department.
    newer = create_brief(db, organization_id=org.id)
    update_brief_fields(db, newer, location_city="Dubai", employment_type="Full-time")

    fields = warm_start_fields(db, org.id)
    # location_city: newest non-empty wins (Dubai over Abu Dhabi).
    assert fields["location_city"] == "Dubai"
    # workplace_type/department only set on the older brief → fall back to it.
    assert fields["workplace_type"] == "Onsite"
    assert fields["department"] == "Eng"
    # employment_type only on the newer brief.
    assert fields["employment_type"] == "Full-time"
    # location_country was never set → omitted entirely (only resolved keys).
    assert "location_country" not in fields


def test_warm_start_fields_empty_when_no_prior_values(db):
    org = _org(db)
    create_brief(db, organization_id=org.id)  # blank brief, nothing to inherit
    assert warm_start_fields(db, org.id) == {}


def test_warm_start_fields_excludes_given_brief(db):
    org = _org(db)
    prior = create_brief(db, organization_id=org.id)
    update_brief_fields(db, prior, location_city="Dubai")
    current = create_brief(db, organization_id=org.id)
    update_brief_fields(db, current, location_city="Riyadh")
    # Excluding ``current`` falls back to the prior brief's value.
    assert warm_start_fields(db, org.id, exclude_brief_id=current.id)["location_city"] == "Dubai"


def test_warm_start_fields_scoped_to_org(db):
    org_a = _org(db)
    other = Organization(name="Other", slug="other")
    db.add(other)
    db.flush()
    foreign = create_brief(db, organization_id=other.id)
    update_brief_fields(db, foreign, location_city="London")
    # A brief in another org must not bleed into org_a's warm-start.
    assert warm_start_fields(db, org_a.id) == {}


def test_recent_role_titles_newest_first_excludes_blank_and_current(db):
    org = _org(db)
    first = create_brief(db, organization_id=org.id)
    update_brief_fields(db, first, title="Backend Engineer")
    second = create_brief(db, organization_id=org.id)
    update_brief_fields(db, second, title="Data Scientist")
    blank = create_brief(db, organization_id=org.id)  # no title → skipped
    current = create_brief(db, organization_id=org.id)
    update_brief_fields(db, current, title="Product Manager")

    titles = recent_role_titles(db, org.id, exclude_brief_id=current.id)
    # Newest-first, blank skipped, current excluded.
    assert titles == ["Data Scientist", "Backend Engineer"]
    assert "Product Manager" not in titles


def test_build_chat_system_prompt_includes_recent_roles_line(db):
    b, _o = _brief(db)
    prompt = build_chat_system_prompt(
        b, resolve_template(_o), focus_gaps=[], recent_titles=["Backend Engineer", "Data Scientist"]
    )
    assert "For context, recent roles at this org: Backend Engineer, Data Scientist." in prompt


def test_build_chat_system_prompt_omits_recent_roles_line_when_none(db):
    b, _o = _brief(db)
    prompt = build_chat_system_prompt(b, resolve_template(_o), focus_gaps=[], recent_titles=[])
    assert "recent roles at this org" not in prompt


def test_build_chat_system_prompt_forbids_unreceipted_publish_claims(db):
    b, org = _brief(db)
    prompt = build_chat_system_prompt(
        b,
        resolve_template(org),
        focus_gaps=[],
        recent_titles=[],
        source_material="[Attached document: role.docx]\nAI Engineer",
        document_turn=True,
    )
    assert "this chat can save brief fields only" in prompt
    assert "NEVER claim" in prompt
    assert "Publish job page" in prompt
    assert "EXTRACT EXHAUSTIVELY" in prompt
    assert "RECOVERABLE SOURCE MATERIAL" in prompt
    assert "Captured-so-far values are authoritative" in prompt
    assert "CHANGE INTENT" in prompt
    assert "change_mode='replace'" in prompt
    assert "add Azure and remove Java" in prompt
    assert "CANONICAL SPEC" in prompt


def test_run_chat_turn_passes_recent_titles_into_system_prompt(db, monkeypatch):
    org = _org(db)
    prior = create_brief(db, organization_id=org.id)
    update_brief_fields(db, prior, title="Staff Engineer")
    current = create_brief(db, organization_id=org.id)
    seed_opening_message(current, resolve_template(org))
    db.flush()

    seen = {}

    def fake_generate_structured(client, **kwargs):
        seen["system"] = kwargs["system"]
        return StructuredResult(value=ChatCapture(assistant_reply="ok"), ok=True)

    monkeypatch.setattr(chat, "generate_structured", fake_generate_structured)
    run_chat_turn(db, current, message="hiring a new role", client=object(), model="m")
    # The prior brief's title surfaces as warm-start context; the current brief
    # (untitled) is excluded so it never lists itself.
    assert "recent roles at this org: Staff Engineer." in seen["system"]


# --------------------------------------------------------------------------- #
# Warm-start from the org's REAL roles (roles.workable_job_data)
# --------------------------------------------------------------------------- #
def _role(db, org, *, workable_job_data=None, created_at=None, deleted_at=None):
    """A Role with optional Workable payload, explicit ``created_at`` (so newest-
    first ordering is deterministic under SQLite), and optional soft-delete."""
    role = Role(
        organization_id=org.id,
        name="Role",
        source="workable",
        workable_job_data=workable_job_data,
        created_at=created_at or datetime.now(timezone.utc),
        deleted_at=deleted_at,
    )
    db.add(role)
    db.flush()
    return role


def test_warm_start_from_roles_normalises_workable_job_data(db):
    org = _org(db)
    _role(
        db, org,
        workable_job_data={
            "title": "Senior Backend Engineer",
            "department": "Engineering",
            "workplace_type": "hybrid",
            "employment_type": "full_time",
            "location": {
                "location_str": "Abu Dhabi, Abu Dhabi, United Arab Emirates",
                "city": "Abu Dhabi",
                "country": "United Arab Emirates",
            },
            "state": "published",
        },
    )
    fields = warm_start_from_roles(db, org.id)
    # snake_case Workable vocab → the template's nice select labels.
    assert fields["workplace_type"] == "Hybrid"
    assert fields["employment_type"] == "Full-time"
    # Structured location used directly.
    assert fields["location_city"] == "Abu Dhabi"
    assert fields["location_country"] == "United Arab Emirates"
    assert fields["department"] == "Engineering"


def test_warm_start_from_roles_recency_biased_per_field(db):
    org = _org(db)
    now = datetime.now(timezone.utc)
    # Older role: has workplace_type + department.
    _role(
        db, org,
        created_at=now - timedelta(days=2),
        workable_job_data={
            "workplace_type": "onsite",
            "department": "Data",
            "location": {"city": "Riyadh", "country": "Saudi Arabia"},
        },
    )
    # Newer role: a different city, remote, but no department.
    _role(
        db, org,
        created_at=now,
        workable_job_data={
            "workplace_type": "remote",
            "location": {"city": "Dubai", "country": "United Arab Emirates"},
        },
    )
    fields = warm_start_from_roles(db, org.id)
    # Newest non-empty per field wins.
    assert fields["workplace_type"] == "Remote"
    assert fields["location_city"] == "Dubai"
    assert fields["location_country"] == "United Arab Emirates"
    # Only the older role had a department → falls back to it.
    assert fields["department"] == "Data"


def test_warm_start_from_roles_falls_back_to_location_str_split(db):
    org = _org(db)
    _role(
        db, org,
        workable_job_data={
            "workplace_type": "Hybrid",  # already a nice label → tolerated
            "employment_type": "contract",
            # No structured city/country — only the human string.
            "location": {"location_str": "London, England, United Kingdom"},
        },
    )
    fields = warm_start_from_roles(db, org.id)
    assert fields["workplace_type"] == "Hybrid"
    assert fields["employment_type"] == "Contract"
    # First part = city, last part = country.
    assert fields["location_city"] == "London"
    assert fields["location_country"] == "United Kingdom"


def test_warm_start_from_roles_skips_missing_and_odd_job_data(db):
    org = _org(db)
    now = datetime.now(timezone.utc)
    # Newest first: a role with no payload, then a non-dict location + null
    # department, then a clean role. None should raise; the clean values win.
    _role(db, org, created_at=now, workable_job_data=None)
    _role(
        db, org,
        created_at=now - timedelta(days=1),
        workable_job_data={
            "workplace_type": "weird-value",   # unrecognised → skipped
            "department": None,                 # null → skipped
            "location": "Just a string city",  # non-dict location
            "employment_type": "part_time",
        },
    )
    _role(
        db, org,
        created_at=now - timedelta(days=2),
        workable_job_data={
            "workplace_type": "remote",
            "location": {"city": "Cairo", "country": "Egypt"},
        },
    )
    fields = warm_start_from_roles(db, org.id)
    # part_time from the odd role (normalised); workplace_type falls through to
    # the clean older role since "weird-value" was unrecognised.
    assert fields["employment_type"] == "Part-time"
    assert fields["workplace_type"] == "Remote"
    # The non-dict location string was treated as the city.
    assert fields["location_city"] == "Just a string city"
    assert fields["location_country"] == "Egypt"  # from the clean older role
    # department never resolved (only null seen) → absent.
    assert "department" not in fields


def test_warm_start_from_roles_excludes_soft_deleted_and_scopes_org(db):
    org = _org(db)
    other = Organization(name="Other", slug="other")
    db.add(other)
    db.flush()
    # A soft-deleted role in-org must be ignored.
    _role(
        db, org,
        deleted_at=datetime.now(timezone.utc),
        workable_job_data={"workplace_type": "onsite"},
    )
    # A role in another org must not bleed in.
    _role(db, other, workable_job_data={"workplace_type": "remote"})
    assert warm_start_from_roles(db, org.id) == {}


def test_warm_start_from_roles_empty_when_no_roles(db):
    org = _org(db)
    assert warm_start_from_roles(db, org.id) == {}


# --------------------------------------------------------------------------- #
# Combined warm-start: brief value WINS, roles fill the rest
# --------------------------------------------------------------------------- #
def test_warm_start_fields_brief_value_wins_over_role(db):
    org = _org(db)
    # A role provides workplace_type + location + employment_type.
    _role(
        db, org,
        workable_job_data={
            "workplace_type": "remote",
            "employment_type": "full_time",
            "location": {"city": "Dubai", "country": "United Arab Emirates"},
        },
    )
    # A recruiter's own recent brief sets workplace_type explicitly.
    brief = create_brief(db, organization_id=org.id)
    update_brief_fields(db, brief, workplace_type="Onsite")

    fields = warm_start_fields(db, org.id)
    # Brief value wins for the field it set...
    assert fields["workplace_type"] == "Onsite"
    # ...and the role fills every field the briefs left empty.
    assert fields["employment_type"] == "Full-time"
    assert fields["location_city"] == "Dubai"
    assert fields["location_country"] == "United Arab Emirates"


def test_warm_start_fields_uses_roles_when_no_briefs(db):
    org = _org(db)
    _role(
        db, org,
        workable_job_data={
            "workplace_type": "hybrid",
            "employment_type": "full_time",
            "location": {"city": "Abu Dhabi", "country": "United Arab Emirates"},
        },
    )
    # No briefs at all → everything comes from the org's real role history.
    fields = warm_start_fields(db, org.id)
    assert fields["workplace_type"] == "Hybrid"
    assert fields["employment_type"] == "Full-time"
    assert fields["location_city"] == "Abu Dhabi"
    assert fields["location_country"] == "United Arab Emirates"


def test_warm_start_fields_empty_with_no_briefs_or_roles(db):
    org = _org(db)
    assert warm_start_fields(db, org.id) == {}


# --------------------------------------------------------------------------- #
# derive_company_blurb — one-time, cached "About the company" extraction
# --------------------------------------------------------------------------- #
def test_derive_company_blurb_extracts_then_caches(db, monkeypatch):
    org = _org(db)
    db.add(Role(
        organization_id=org.id, name="Engineer", source="workable",
        job_spec_text="About Acme\nWe build payroll software.\n\nRequirements\n- Python",
    ))
    db.flush()
    calls = {"n": 0}

    def fake_generate_structured(client, **kwargs):
        calls["n"] += 1
        return StructuredResult(
            value=chat.CompanyBlurbDraft(company_description="Acme builds payroll software."),
            ok=True,
        )

    monkeypatch.setattr(chat, "generate_structured", fake_generate_structured)
    blurb = chat.derive_company_blurb(db, org.id, client=object())
    assert blurb == "Acme builds payroll software."
    assert org.company_blurb == "Acme builds payroll software."
    # Cached — a second call returns the stored blurb WITHOUT another LLM call.
    assert chat.derive_company_blurb(db, org.id, client=object()) == "Acme builds payroll software."
    assert calls["n"] == 1


def test_derive_company_blurb_no_specs_caches_empty_and_skips_llm(db, monkeypatch):
    org = _org(db)
    calls = {"n": 0}

    def fake_generate_structured(client, **kwargs):
        calls["n"] += 1
        return StructuredResult(value=chat.CompanyBlurbDraft(company_description="x"), ok=True)

    monkeypatch.setattr(chat, "generate_structured", fake_generate_structured)
    assert chat.derive_company_blurb(db, org.id, client=object()) is None
    assert org.company_blurb == ""   # no-result cached so we don't re-call
    assert calls["n"] == 0           # never invoked the model (no specs to read)
